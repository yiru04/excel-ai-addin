"""
Vercel Python Serverless Function：接收 PDF / Word 檔案（base64），
用伺服器端的成熟函式庫解析成純文字後回傳。

這個端點取代了原本在瀏覽器裡用 pdf.js / mammoth.js 做解析的做法——
跟 Claude、ChatGPT 平常處理檔案的方式一樣，解析工作在伺服器端做，
不會受瀏覽器端 CDN 載入、ES Module 相容性這些問題影響。
"""

from http.server import BaseHTTPRequestHandler
import json
import base64
import io

MAX_PDF_PAGES = 30


def extract_pdf_text(file_bytes: bytes) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        total_pages = len(pdf.pages)
        max_pages = min(total_pages, MAX_PDF_PAGES)
        for i in range(max_pages):
            page_text = pdf.pages[i].extract_text() or ""
            text_parts.append(page_text)

        note = ""
        if total_pages > max_pages:
            note = f"\n（僅擷取前 {max_pages} 頁，PDF 共 {total_pages} 頁）"

    return "\n\n".join(text_parts) + note


def extract_docx_text(file_bytes: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # 表格內容也一併擷取，不然報表類的 Word 檔會漏掉大部分資訊
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n".join(parts)


class handler(BaseHTTPRequestHandler):
    def _set_cors_headers(self):
        # 只允許我們自己的 GitHub Pages 網域呼叫，不是完全開放給任何網站使用
        self.send_header('Access-Control-Allow-Origin', 'https://yiru04.github.io')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')

    def do_OPTIONS(self):
        self.send_response(204)
        self._set_cors_headers()
        self.end_headers()

    def do_POST(self):
        try:
            content_length = int(self.headers.get('Content-Length', 0))
            if content_length == 0 or content_length > 15 * 1024 * 1024:
                raise ValueError('檔案大小不符合限制（上限約 15MB）')

            body = self.rfile.read(content_length)
            data = json.loads(body)

            file_type = data.get('type')
            file_b64 = data.get('content', '')
            if not file_b64:
                raise ValueError('沒有收到檔案內容')

            file_bytes = base64.b64decode(file_b64)

            if file_type == 'pdf':
                text = extract_pdf_text(file_bytes)
            elif file_type == 'docx':
                text = extract_docx_text(file_bytes)
            else:
                raise ValueError(f'不支援的檔案類型：{file_type}')

            self.send_response(200)
            self._set_cors_headers()
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.end_headers()
            self.wfile.write(json.dumps({'text': text}, ensure_ascii=False).encode('utf-8'))

        except Exception as e:
            self.send_response(500)
            self._set_cors_headers()
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.end_headers()
            self.wfile.write(json.dumps({'error': str(e)}, ensure_ascii=False).encode('utf-8'))
